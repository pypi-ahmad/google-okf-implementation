from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from ingest.parser import DocumentParseError, DocumentParser


def test_parse_markdown_extracts_structure_tables_and_section_chunks(tmp_path: Path) -> None:
    path = tmp_path / "runbook.md"
    path.write_text(
        "\n".join(
            [
                "# Payments",
                "",
                "PAYMENT_ALPHA PAYMENT_ALPHA PAYMENT_ALPHA PAYMENT_ALPHA PAYMENT_ALPHA PAYMENT_ALPHA PAYMENT_ALPHA",
                "",
                "| metric | value |",
                "| --- | --- |",
                "| payment_failures | 42 |",
                "",
                "## Orders",
                "",
                "ORDERS_BETA ORDERS_BETA ORDERS_BETA ORDERS_BETA ORDERS_BETA ORDERS_BETA ORDERS_BETA",
            ]
        ),
        encoding="utf-8",
    )

    parser = DocumentParser(chunk_size_chars=120, chunk_overlap_chars=20)
    parsed = parser.parse(path)

    assert parsed.file_type == "markdown"
    assert parsed.headings[0].text == "Payments"
    assert parsed.headings[1].text == "Orders"
    assert len(parsed.tables) == 1
    assert parsed.tables[0].headers == ["metric", "value"]
    assert parsed.tables[0].rows[0] == ["payment_failures", "42"]

    payment_chunks = [chunk for chunk in parsed.chunks if chunk.section_id == "s1"]
    order_chunks = [chunk for chunk in parsed.chunks if chunk.section_id == "s2"]

    assert payment_chunks
    assert order_chunks
    assert all("ORDERS_BETA" not in chunk.content for chunk in payment_chunks)
    assert all("PAYMENT_ALPHA" not in chunk.content for chunk in order_chunks)


def test_parse_docx_extracts_headings_tables_and_metadata(tmp_path: Path) -> None:
    path = tmp_path / "playbook.docx"

    doc = DocxDocument()
    doc.add_heading("Payment Incident", level=1)
    doc.add_paragraph("Escalate to SRE on PagerDuty.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "owner"
    table.rows[0].cells[1].text = "team"
    table.rows[1].cells[0].text = "payments-api"
    table.rows[1].cells[1].text = "sre"
    doc.core_properties.author = "SRE Team"
    doc.core_properties.created = datetime(2026, 7, 3, tzinfo=timezone.utc)
    doc.save(path)

    parsed = DocumentParser().parse(path)

    assert parsed.file_type == "docx"
    assert parsed.headings[0].text == "Payment Incident"
    assert "Escalate to SRE" in parsed.content
    assert parsed.metadata["author"] == "SRE Team"
    assert parsed.metadata["creation_date"] is not None
    assert len(parsed.tables) == 1
    assert parsed.tables[0].headers == ["owner", "team"]


def test_parse_pdf_extracts_metadata_and_page_sections(tmp_path: Path) -> None:
    path = tmp_path / "architecture.pdf"

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_blank_page(width=200, height=200)
    writer.add_metadata({"/Author": "Data Platform", "/CreationDate": "D:20260703083000Z"})
    with path.open("wb") as stream:
        writer.write(stream)

    parsed = DocumentParser().parse(path)

    assert parsed.file_type == "pdf"
    assert parsed.metadata["author"] == "Data Platform"
    assert parsed.metadata["creation_date"] is not None
    assert len(parsed.sections) == 2
    assert parsed.headings[0].text == "Page 1"


def test_parse_csv_extracts_table_and_metadata(tmp_path: Path) -> None:
    path = tmp_path / "datasets.csv"
    pd.DataFrame(
        [
            {"dataset": "orders", "owner": "analytics"},
            {"dataset": "payments", "owner": "finance"},
        ]
    ).to_csv(path, index=False)

    parsed = DocumentParser().parse(path)

    assert parsed.file_type == "csv"
    assert len(parsed.tables) == 1
    assert parsed.tables[0].headers == ["dataset", "owner"]
    assert parsed.metadata["row_count"] == "2"
    assert parsed.metadata["column_count"] == "2"


def test_parse_html_extracts_headings_tables_and_author(tmp_path: Path) -> None:
    path = tmp_path / "incident.html"
    path.write_text(
        """
        <html>
          <head>
            <meta name="author" content="Ops Team" />
            <meta property="article:published_time" content="2026-07-01" />
          </head>
          <body>
            <h1>Incident</h1>
            <p>Payment API latency spike.</p>
            <table>
              <tr><th>service</th><th>status</th></tr>
              <tr><td>payments-api</td><td>degraded</td></tr>
            </table>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    parsed = DocumentParser().parse(path)

    assert parsed.file_type == "html"
    assert parsed.headings[0].text == "Incident"
    assert "Payment API latency spike" in parsed.content
    assert parsed.metadata["author"] == "Ops Team"
    assert parsed.metadata["creation_date"] is not None
    assert len(parsed.tables) == 1
    assert parsed.tables[0].headers == ["service", "status"]


def test_parse_malformed_pdf_recovers_with_errors(tmp_path: Path) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"this is not a valid pdf")

    parsed = DocumentParser(recover_errors=True).parse(path)

    assert parsed.file_type == "pdf"
    assert parsed.errors
    assert parsed.content == ""
    assert parsed.chunks == []


def test_parse_malformed_pdf_raises_when_recovery_disabled(tmp_path: Path) -> None:
    path = tmp_path / "broken-strict.pdf"
    path.write_bytes(b"this is not a valid pdf")

    with pytest.raises(DocumentParseError):
        DocumentParser(recover_errors=False).parse(path)
