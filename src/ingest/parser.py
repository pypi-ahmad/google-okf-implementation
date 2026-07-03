"""Robust multi-format enterprise document parser with normalization and chunking."""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from hashlib import sha256
from pathlib import Path
from typing import Any

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from loguru import logger
from pypdf import PdfReader

SUPPORTED_FILE_TYPES: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "markdown",
    ".csv": "csv",
    ".html": "html",
    ".htm": "html",
}

_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MD_TABLE_SEPARATOR_PATTERN = re.compile(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$")


class DocumentParseError(RuntimeError):
    """Raised when parsing fails and recovery is disabled."""


@dataclass(slots=True)
class DocumentHeading:
    """Structured heading extracted from a source document."""

    level: int
    text: str
    section_id: str
    line_number: int | None = None


@dataclass(slots=True)
class DocumentTable:
    """Normalized table representation."""

    headers: list[str]
    rows: list[list[str]]
    source: str
    section_id: str = "root"


@dataclass(slots=True)
class DocumentSection:
    """Semantic section used for chunking and downstream extraction."""

    section_id: str
    title: str
    level: int
    content: str


@dataclass(slots=True)
class DocumentChunk:
    """Chunk generated from one semantic section."""

    chunk_id: str
    section_id: str
    content: str
    index: int


@dataclass(slots=True)
class ParsedDocument:
    """Normalized parse output with content, metadata, structure, and provenance."""

    file_path: Path
    file_type: str
    content: str
    metadata: dict[str, str | None]
    provenance: dict[str, str | int]
    headings: list[DocumentHeading] = field(default_factory=list)
    tables: list[DocumentTable] = field(default_factory=list)
    sections: list[DocumentSection] = field(default_factory=list)
    chunks: list[DocumentChunk] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        """Serialize to a JSON-friendly payload."""

        return {
            "file_path": self.file_path.as_posix(),
            "file_type": self.file_type,
            "content": self.content,
            "metadata": self.metadata,
            "provenance": self.provenance,
            "headings": [
                {
                    "level": heading.level,
                    "text": heading.text,
                    "section_id": heading.section_id,
                    "line_number": heading.line_number,
                }
                for heading in self.headings
            ],
            "tables": [
                {
                    "headers": table.headers,
                    "rows": table.rows,
                    "source": table.source,
                    "section_id": table.section_id,
                }
                for table in self.tables
            ],
            "sections": [
                {
                    "section_id": section.section_id,
                    "title": section.title,
                    "level": section.level,
                    "content": section.content,
                }
                for section in self.sections
            ],
            "chunks": [
                {
                    "chunk_id": chunk.chunk_id,
                    "section_id": chunk.section_id,
                    "content": chunk.content,
                    "index": chunk.index,
                }
                for chunk in self.chunks
            ],
            "errors": self.errors,
        }


@dataclass(slots=True)
class _LoaderResult:
    """Internal loader response before parser-level post-processing."""

    content: str
    headings: list[DocumentHeading]
    tables: list[DocumentTable]
    sections: list[DocumentSection]
    metadata: dict[str, str | None] = field(default_factory=dict)


class DocumentParser:
    """Parse enterprise documentation from common formats.

    The parser normalizes content to a common structure with headings, sections,
    tables, provenance metadata, and deterministic chunks.
    """

    def __init__(
        self,
        chunk_size_chars: int = 1200,
        chunk_overlap_chars: int = 150,
        recover_errors: bool = True,
    ) -> None:
        if chunk_size_chars <= 0:
            raise ValueError("chunk_size_chars must be positive")
        if chunk_overlap_chars < 0 or chunk_overlap_chars >= chunk_size_chars:
            raise ValueError("chunk_overlap_chars must be >= 0 and < chunk_size_chars")

        self.chunk_size_chars = chunk_size_chars
        self.chunk_overlap_chars = chunk_overlap_chars
        self.recover_errors = recover_errors

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """Parse one file into a unified normalized document object."""

        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"File not found: {path}")

        suffix = path.suffix.lower()
        file_type = SUPPORTED_FILE_TYPES.get(suffix)
        if file_type is None:
            raise ValueError(f"Unsupported file extension: {suffix}")

        logger.debug("Parsing document: {}", path)

        loader_result: _LoaderResult
        errors: list[str] = []
        try:
            loader_result = self._resolve_loader(suffix)(path)
        except Exception as exc:  # noqa: BLE001
            message = f"{exc.__class__.__name__}: {exc}"
            logger.exception("Failed parsing {}", path)
            if not self.recover_errors:
                raise DocumentParseError(message) from exc

            loader_result = _LoaderResult(content="", headings=[], tables=[], sections=[])
            errors.append(message)

        metadata = self._base_metadata(
            path,
            author=loader_result.metadata.get("author"),
            creation_date=loader_result.metadata.get("creation_date"),
        )

        # Preserve additional document-level metadata provided by loaders.
        for key, value in loader_result.metadata.items():
            if key not in {"author", "creation_date"}:
                metadata[key] = value

        sections = loader_result.sections if loader_result.sections else self._fallback_sections(loader_result.content)
        chunks = self._chunk_sections(path, sections)

        return ParsedDocument(
            file_path=path,
            file_type=file_type,
            content=loader_result.content,
            metadata=metadata,
            provenance=self._build_provenance(path, file_type),
            headings=loader_result.headings,
            tables=loader_result.tables,
            sections=sections,
            chunks=chunks,
            errors=errors,
        )

    def _resolve_loader(self, suffix: str):
        if suffix == ".pdf":
            return self._parse_pdf
        if suffix == ".docx":
            return self._parse_docx
        if suffix in {".md", ".markdown", ".txt"}:
            return self._parse_markdown
        if suffix == ".csv":
            return self._parse_csv
        if suffix in {".html", ".htm"}:
            return self._parse_html

        raise ValueError(f"Unsupported document extension: {suffix}")

    def _parse_markdown(self, path: Path) -> _LoaderResult:
        text = path.read_text(encoding="utf-8", errors="ignore").replace("\r\n", "\n").replace("\r", "\n").strip()
        headings, sections = self._split_markdown_sections(text)
        tables = self._extract_markdown_tables(text)
        return _LoaderResult(content=text, headings=headings, tables=tables, sections=sections)

    def _parse_pdf(self, path: Path) -> _LoaderResult:
        reader = PdfReader(str(path))
        metadata_raw: Any = reader.metadata or {}

        author = self._clean_string(getattr(metadata_raw, "author", None))
        creation_date = self._normalize_datetime(getattr(metadata_raw, "creation_date", None))

        if isinstance(metadata_raw, dict):
            author = author or self._clean_string(metadata_raw.get("/Author") or metadata_raw.get("Author"))
            creation_date = creation_date or self._normalize_datetime(
                metadata_raw.get("/CreationDate") or metadata_raw.get("CreationDate")
            )

        headings: list[DocumentHeading] = []
        sections: list[DocumentSection] = []

        for page_index, page in enumerate(reader.pages, start=1):
            title = f"Page {page_index}"
            section_id = f"page-{page_index}"
            page_text = (page.extract_text() or "").strip()
            section_content = f"# {title}\n{page_text}".strip()

            headings.append(DocumentHeading(level=1, text=title, section_id=section_id))
            sections.append(
                DocumentSection(
                    section_id=section_id,
                    title=title,
                    level=1,
                    content=section_content,
                )
            )

        content = "\n\n".join(section.content for section in sections if section.content).strip()

        return _LoaderResult(
            content=content,
            headings=headings,
            tables=[],
            sections=sections,
            metadata={"author": author, "creation_date": creation_date},
        )

    def _parse_docx(self, path: Path) -> _LoaderResult:
        doc = DocxDocument(str(path))

        blocks: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style is not None else ""
            heading_level = self._docx_heading_level(style_name)
            if heading_level is not None:
                blocks.append(f"{'#' * heading_level} {text}")
            else:
                blocks.append(text)

        content = "\n\n".join(blocks).strip()
        headings, sections = self._split_markdown_sections(content)
        tables = self._extract_docx_tables(doc)

        props = doc.core_properties
        author = self._clean_string(props.author)
        creation_date = self._normalize_datetime(props.created)

        return _LoaderResult(
            content=content,
            headings=headings,
            tables=tables,
            sections=sections,
            metadata={"author": author, "creation_date": creation_date},
        )

    def _parse_csv(self, path: Path) -> _LoaderResult:
        frame = pd.read_csv(
            path,
            dtype=str,
            keep_default_na=False,
            engine="python",
            on_bad_lines="skip",
            encoding_errors="ignore",
        ).fillna("")

        headers = [str(column).strip() for column in frame.columns.tolist()]
        rows = [[str(value) for value in row] for row in frame.astype(str).values.tolist()]

        table = DocumentTable(headers=headers, rows=rows, source="csv", section_id="root")
        content = frame.to_csv(index=False, lineterminator="\n").strip()
        sections = [DocumentSection(section_id="root", title=path.stem, level=0, content=content)] if content else []

        return _LoaderResult(
            content=content,
            headings=[],
            tables=[table],
            sections=sections,
            metadata={
                "row_count": str(len(rows)),
                "column_count": str(len(headers)),
            },
        )

    def _parse_html(self, path: Path) -> _LoaderResult:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        author_tag = soup.find("meta", attrs={"name": "author"})
        created_tag = (
            soup.find("meta", attrs={"name": "date"})
            or soup.find("meta", attrs={"name": "published"})
            or soup.find("meta", attrs={"property": "article:published_time"})
        )

        content_blocks: list[str] = []
        body = soup.body or soup
        for element in body.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre"]):
            text = element.get_text(" ", strip=True)
            if not text:
                continue

            if element.name and re.fullmatch(r"h[1-6]", element.name):
                level = int(element.name[1])
                content_blocks.append(f"{'#' * level} {text}")
            else:
                content_blocks.append(text)

        tables = self._extract_html_tables(soup)
        content_blocks.extend(self._table_to_text(table) for table in tables)

        content = "\n\n".join(content_blocks).strip()
        headings, sections = self._split_markdown_sections(content)

        return _LoaderResult(
            content=content,
            headings=headings,
            tables=tables,
            sections=sections,
            metadata={
                "author": self._clean_string(author_tag.get("content") if author_tag else None),
                "creation_date": self._normalize_datetime(created_tag.get("content") if created_tag else None),
            },
        )

    def _split_markdown_sections(self, content: str) -> tuple[list[DocumentHeading], list[DocumentSection]]:
        if not content:
            return [], []

        lines = content.splitlines()
        heading_markers: list[tuple[int, int, str, str]] = []
        headings: list[DocumentHeading] = []

        for line_index, line in enumerate(lines, start=1):
            match = _HEADING_PATTERN.match(line.strip())
            if not match:
                continue

            level = len(match.group(1))
            title = match.group(2).strip()
            section_id = f"s{len(headings) + 1}"
            heading_markers.append((line_index, level, title, section_id))
            headings.append(DocumentHeading(level=level, text=title, section_id=section_id, line_number=line_index))

        if not heading_markers:
            root_content = content.strip()
            if not root_content:
                return headings, []
            return headings, [DocumentSection(section_id="root", title="root", level=0, content=root_content)]

        sections: list[DocumentSection] = []

        first_heading_line = heading_markers[0][0]
        preamble = "\n".join(lines[: first_heading_line - 1]).strip()
        if preamble:
            sections.append(DocumentSection(section_id="root", title="root", level=0, content=preamble))

        for index, marker in enumerate(heading_markers):
            start_line, level, title, section_id = marker
            end_line = heading_markers[index + 1][0] - 1 if index + 1 < len(heading_markers) else len(lines)
            section_text = "\n".join(lines[start_line - 1 : end_line]).strip()
            if section_text:
                sections.append(DocumentSection(section_id=section_id, title=title, level=level, content=section_text))

        return headings, sections

    def _extract_markdown_tables(self, content: str) -> list[DocumentTable]:
        lines = content.splitlines()
        tables: list[DocumentTable] = []

        idx = 0
        while idx < len(lines) - 1:
            header_line = lines[idx].strip()
            separator_line = lines[idx + 1].strip()

            if "|" not in header_line or not _MD_TABLE_SEPARATOR_PATTERN.match(separator_line):
                idx += 1
                continue

            headers = self._split_markdown_row(header_line)
            if not headers:
                idx += 1
                continue

            idx += 2
            rows: list[list[str]] = []
            while idx < len(lines):
                row_line = lines[idx].strip()
                if not row_line or "|" not in row_line:
                    break
                cells = self._split_markdown_row(row_line)
                if cells:
                    rows.append(self._normalize_row(cells, len(headers)))
                idx += 1

            tables.append(DocumentTable(headers=headers, rows=rows, source="markdown", section_id="root"))

        return tables

    def _extract_docx_tables(self, doc: DocxDocument) -> list[DocumentTable]:
        tables: list[DocumentTable] = []

        for table in doc.tables:
            matrix: list[list[str]] = []
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    matrix.append(values)

            if not matrix:
                continue

            headers = matrix[0]
            if not headers:
                continue

            rows = [self._normalize_row(row, len(headers)) for row in matrix[1:]]
            tables.append(DocumentTable(headers=headers, rows=rows, source="docx", section_id="root"))

        return tables

    def _extract_html_tables(self, soup: BeautifulSoup) -> list[DocumentTable]:
        tables: list[DocumentTable] = []

        for table in soup.find_all("table"):
            matrix: list[list[str]] = []
            for row in table.find_all("tr"):
                values = [cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"])]
                if any(values):
                    matrix.append(values)

            if not matrix:
                continue

            headers = matrix[0]
            if not headers:
                continue

            rows = [self._normalize_row(row, len(headers)) for row in matrix[1:]]
            tables.append(DocumentTable(headers=headers, rows=rows, source="html", section_id="root"))

        return tables

    def _table_to_text(self, table: DocumentTable) -> str:
        lines = [", ".join(table.headers)]
        lines.extend(", ".join(self._normalize_row(row, len(table.headers))) for row in table.rows)
        return "\n".join(lines)

    def _split_markdown_row(self, row: str) -> list[str]:
        normalized = row.strip().strip("|")
        if not normalized:
            return []
        return [cell.strip() for cell in normalized.split("|")]

    def _chunk_sections(self, path: Path, sections: list[DocumentSection]) -> list[DocumentChunk]:
        chunks: list[DocumentChunk] = []

        for section in sections:
            for local_index, chunk_text in enumerate(self._chunk_text(section.content)):
                digest = sha256(f"{path.as_posix()}::{section.section_id}::{local_index}::{chunk_text}".encode()).hexdigest()
                chunks.append(
                    DocumentChunk(
                        chunk_id=digest,
                        section_id=section.section_id,
                        content=chunk_text,
                        index=len(chunks),
                    )
                )

        return chunks

    def _chunk_text(self, text: str) -> list[str]:
        normalized = text.strip()
        if not normalized:
            return []

        if len(normalized) <= self.chunk_size_chars:
            return [normalized]

        step = self.chunk_size_chars - self.chunk_overlap_chars
        chunks: list[str] = []
        start = 0

        while start < len(normalized):
            end = min(start + self.chunk_size_chars, len(normalized))
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= len(normalized):
                break
            start += step

        return chunks

    def _fallback_sections(self, content: str) -> list[DocumentSection]:
        text = content.strip()
        if not text:
            return []
        return [DocumentSection(section_id="root", title="root", level=0, content=text)]

    def _build_provenance(self, path: Path, file_type: str) -> dict[str, str | int]:
        stat = path.stat()
        return {
            "source": path.as_posix(),
            "absolute_path": path.resolve().as_posix(),
            "file_type": file_type,
            "size_bytes": stat.st_size,
            "sha256": self._sha256_file(path),
        }

    def _base_metadata(self, path: Path, author: str | None = None, creation_date: str | None = None) -> dict[str, str | None]:
        stat = path.stat()
        fallback_created = datetime.fromtimestamp(stat.st_ctime, tz=timezone.utc).isoformat()
        return {
            "source": str(path),
            "file_name": path.name,
            "extension": path.suffix.lower(),
            "author": author,
            "creation_date": creation_date or fallback_created,
            "modified_date": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
        }

    def _sha256_file(self, path: Path) -> str:
        hasher = sha256()
        with path.open("rb") as stream:
            for chunk in iter(lambda: stream.read(8192), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    def _normalize_datetime(self, value: object) -> str | None:
        if value is None:
            return None

        if isinstance(value, datetime):
            if value.tzinfo is None:
                return value.replace(tzinfo=timezone.utc).isoformat()
            return value.astimezone(timezone.utc).isoformat()

        text = self._clean_string(value)
        if not text:
            return None

        if text.startswith("D:"):
            digits = text[2:16]
            try:
                parsed = datetime.strptime(digits, "%Y%m%d%H%M%S")
                return parsed.replace(tzinfo=timezone.utc).isoformat()
            except ValueError:
                return None

        try:
            parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
            if parsed.tzinfo is None:
                parsed = parsed.replace(tzinfo=timezone.utc)
            return parsed.astimezone(timezone.utc).isoformat()
        except ValueError:
            pass

        for fmt in ["%Y-%m-%d", "%Y/%m/%d"]:
            try:
                parsed = datetime.strptime(text, fmt)
                return parsed.replace(tzinfo=timezone.utc).isoformat()
            except ValueError:
                continue

        return text

    def _clean_string(self, value: object) -> str | None:
        if value is None:
            return None
        text = str(value).strip()
        return text if text else None

    def _normalize_row(self, row: list[str], width: int) -> list[str]:
        normalized = [cell.strip() for cell in row]
        if len(normalized) < width:
            normalized.extend([""] * (width - len(normalized)))
        elif len(normalized) > width:
            normalized = normalized[:width]
        return normalized

    def _docx_heading_level(self, style_name: str) -> int | None:
        match = re.search(r"heading\s*([1-6])", style_name, flags=re.IGNORECASE)
        if not match:
            return None
        return int(match.group(1))
