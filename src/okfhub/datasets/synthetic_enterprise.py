"""Synthetic but realistic enterprise corpus generator."""

from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from fpdf import FPDF

from okfhub.utils.filesystem import ensure_dir


class SyntheticEnterpriseCorpus:
    """Generate a licensing-safe enterprise knowledge corpus.

    Example:
        >>> corpus = SyntheticEnterpriseCorpus(Path("data/raw"))
        >>> corpus.generate()
    """

    def __init__(self, root: Path):
        self._root = root

    def generate(self) -> list[Path]:
        """Generate multi-format documents used for end-to-end demo runs."""

        ensure_dir(self._root)
        outputs: list[Path] = []

        outputs.append(self._write_markdown())
        outputs.append(self._write_runbook())
        outputs.append(self._write_sql_doc())
        outputs.append(self._write_csv())
        outputs.append(self._write_html())
        outputs.append(self._write_docx())
        outputs.append(self._write_pdf())

        return outputs

    def _write_markdown(self) -> Path:
        path = self._root / "apis" / "orders_api.md"
        ensure_dir(path.parent)
        path.write_text(
            """
# Orders API

## Endpoint
`PATCH /v2/orders/{order_id}` updates customer order status and fulfillment metadata.

## Owner
Order Platform Team

## Dependencies
- customer_profile dataset
- orders table
- order_status_sync playbook

## Notes
This endpoint writes to warehouse table `mart.orders` and emits event `order.updated`.
""".strip()
            + "\n",
            encoding="utf-8",
        )
        return path

    def _write_runbook(self) -> Path:
        path = self._root / "runbooks" / "payment_failure_runbook.md"
        ensure_dir(path.parent)
        path.write_text(
            """
# Payment Failure Runbook

## Trigger
Error rate > 3% on Payment API for 5 minutes.

## Steps
1. Check Grafana dashboard `payments-health`.
2. Validate DB connectivity from payment-service.
3. Roll back payment-gateway adapter to last known stable build.

## Owner
Payments SRE

## Related Metrics
- Payment Success Rate
- Failed Transactions
""".strip()
            + "\n",
            encoding="utf-8",
        )
        return path

    def _write_sql_doc(self) -> Path:
        path = self._root / "sql" / "mau_definition.md"
        ensure_dir(path.parent)
        path.write_text(
            """
# Monthly Active Users

MAU is computed as distinct `user_id` where event_time falls in calendar month and event_type in (session_start, purchase, support_ticket).

```sql
SELECT DATE_TRUNC('month', event_time) AS month,
       COUNT(DISTINCT user_id) AS monthly_active_users
FROM analytics.user_events
WHERE event_type IN ('session_start', 'purchase', 'support_ticket')
GROUP BY 1;
```

Owner: Growth Analytics
""".strip()
            + "\n",
            encoding="utf-8",
        )
        return path

    def _write_csv(self) -> Path:
        path = self._root / "data_dictionary" / "orders_table.csv"
        ensure_dir(path.parent)
        df = pd.DataFrame(
            [
                {
                    "column": "order_id",
                    "type": "STRING",
                    "description": "Unique order identifier",
                    "owner": "Data Platform",
                },
                {
                    "column": "customer_id",
                    "type": "STRING",
                    "description": "Customer foreign key",
                    "owner": "Data Platform",
                },
                {
                    "column": "order_status",
                    "type": "STRING",
                    "description": "Current fulfillment status",
                    "owner": "Order Platform",
                },
            ]
        )
        df.to_csv(path, index=False)
        return path

    def _write_html(self) -> Path:
        path = self._root / "incidents" / "payment_outage_report.html"
        ensure_dir(path.parent)
        path.write_text(
            """
<html>
  <body>
    <h1>Incident 2026-05-17: Payment Outage</h1>
    <p>Impact: 28 minutes, 12% checkout failures.</p>
    <h2>Root Cause</h2>
    <p>Connection pool exhaustion in payment-service after unbounded retry loop.</p>
    <h2>Action Items</h2>
    <ul>
      <li>Add circuit breaker for external gateway timeout</li>
      <li>Set retry budget at 2 attempts</li>
      <li>Create canary alert for checkout latency p95 > 900ms</li>
    </ul>
  </body>
</html>
""".strip()
            + "\n",
            encoding="utf-8",
        )
        return path

    def _write_docx(self) -> Path:
        path = self._root / "handbook" / "employee_handbook.docx"
        ensure_dir(path.parent)
        document = DocxDocument()
        document.add_heading("Employee Handbook", level=1)
        document.add_paragraph("All production changes require ticket linkage and reviewer approval.")
        document.add_paragraph("Critical incident severity matrix is defined by customer impact and duration.")
        document.add_paragraph("Knowledge owners must refresh metrics and runbooks every 90 days.")
        document.save(path)
        return path

    def _write_pdf(self) -> Path:
        path = self._root / "architecture" / "data_platform_overview.pdf"
        ensure_dir(path.parent)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        lines = [
            "Data Platform Overview",
            "",
            "Core datasets: customer_profile, orders, payments, invoices.",
            "Primary warehouse: BigQuery in analytics-prod project.",
            "SLA: dashboard freshness < 30 minutes for executive metrics.",
            "Owner: Data Platform Engineering",
        ]
        for line in lines:
            if not line:
                pdf.ln(4)
                continue
            pdf.multi_cell(w=180, h=8, text=line)

        pdf.output(str(path))
        return path
